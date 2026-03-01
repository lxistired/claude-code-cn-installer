"""生成 Claude Code 安装工具使用说明 (Word + PDF)"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from fpdf import FPDF

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ── 内容定义 ──────────────────────────────────────────────

TITLE = "Claude Code 一键安装工具 使用说明"
SUBTITLE = "Windows 中国版 — 智谱 GLM"

SECTIONS = [
    {
        "heading": "一、工具简介",
        "paragraphs": [
            "本工具帮助 Windows 用户一键安装 Claude Code，并自动配置智谱 GLM API，无需手动安装 Node.js、Git 等依赖。",
        ],
    },
    {
        "heading": "二、系统要求",
        "bullets": [
            "Windows 10 或更高版本",
            "需要管理员权限（安装时会自动请求）",
            "需要网络连接（下载依赖包）",
        ],
    },
    {
        "heading": "三、安装步骤",
        "numbered": [
            "将本工具的所有文件下载到同一个文件夹中",
            '右键点击「一键安装.bat」-> 选择「以管理员身份运行」',
            "安装程序会自动完成以下工作：",
            "  - 安装 Node.js（从国内镜像下载）",
            "  - 安装 Git",
            "  - 配置 npm 国内镜像源",
            "  - 安装 Claude Code",
            "安装完成后，按照提示选择 GLM 模型、输入 API Key",
        ],
    },
    {
        "heading": "四、支持的智谱 GLM 模型",
        "paragraphs": [
            "本工具通过智谱 GLM 的 Anthropic 兼容端点接入，注册地址：open.bigmodel.cn",
        ],
        "table": {
            "headers": ["模型", "说明"],
            "rows": [
                ["glm-5", "旗舰 745B MoE（对标 Opus）推荐"],
                ["glm-4.7", "编程增强（对标 Sonnet）"],
                ["glm-4.5", "Agent 基座，工具调用优化"],
                ["glm-4.7-flash", "30B MoE 轻量快速（对标 Haiku）"],
                ["glm-4-flash", "免费模型，轻量任务"],
                ["glm-4.5-air", "轻量快速，低成本"],
            ],
        },
        "paragraphs_after": [
            "请先到智谱开放平台注册账号并获取 API Key（右上角头像 → API 密钥 → 创建）。",
        ],
    },
    {
        "heading": "五、安装后使用",
        "paragraphs": [
            "安装完成后，请关闭当前窗口，打开一个新的 PowerShell 或 CMD 窗口，输入以下命令：",
        ],
        "code": [
            "claude             (启动交互模式)",
            "claude --help      (查看帮助)",
            "claude --version   (查看版本)",
        ],
    },
    {
        "heading": "六、更换 API Key / 切换模型",
        "paragraphs": [
            '如需更换 API Key 或切换模型，双击运行「配置API.bat」即可。',
        ],
        "bullets": [
            "查看当前配置（含当前使用的模型名称）",
            "重新选择 GLM 模型",
            "更换 API Key",
            "测试 API 连接",
        ],
    },
    {
        "heading": "七、常见问题",
        "qa": [
            ('输入 claude 提示"不是内部或外部命令"？',
             "关闭当前终端，打开新的 PowerShell 或 CMD 窗口再试。"),
            ("安装速度很慢？",
             "脚本已自动配置国内镜像源，如仍慢请检查网络连接。"),
            ("API 连接测试失败？",
             "请检查 API Key 是否正确、账户是否有余额、网络是否正常。"),
        ],
    },
    {
        "heading": "八、文件清单",
        "table": {
            "headers": ["文件", "用途"],
            "rows": [
                ["一键安装.bat", "安装入口，双击运行"],
                ["配置API.bat", "API 配置工具，双击运行"],
                ["install.ps1", "安装主脚本（由 bat 调用）"],
                ["configure-api.ps1", "API 配置脚本（由 bat 调用）"],
            ],
        },
    },
]

FOOTER = "仅供学习和个人使用。Claude Code 版权归 Anthropic 所有。"


# ── Word 生成 ─────────────────────────────────────────────

def make_docx(path):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(11)

    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(TITLE)
    run.bold = True
    run.font.size = Pt(20)
    run.font.name = "微软雅黑"
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(SUBTITLE)
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    for sec in SECTIONS:
        h = doc.add_heading(sec["heading"], level=2)
        for run in h.runs:
            run.font.name = "微软雅黑"
            run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)

        for text in sec.get("paragraphs", []):
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)

        for item in sec.get("bullets", []):
            doc.add_paragraph(item, style="List Bullet")

        for i, item in enumerate(sec.get("numbered", []), 1):
            if item.startswith("  -"):
                # 子项，作为普通缩进文本
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(2)
                p.add_run(item.strip())
            else:
                p = doc.add_paragraph(style="List Number")
                p.add_run(item)

        tbl = sec.get("table")
        if tbl:
            headers = tbl["headers"]
            rows = tbl["rows"]
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = "Light Shading Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for j, h_text in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = h_text
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            for i, row in enumerate(rows):
                for j, val in enumerate(row):
                    table.rows[i + 1].cells[j].text = val
            doc.add_paragraph()

        for text in sec.get("paragraphs_after", []):
            p = doc.add_paragraph(text)
            p.paragraph_format.space_after = Pt(6)

        code_lines = sec.get("code")
        if code_lines:
            for line in code_lines:
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.font.name = "Consolas"
                run.font.size = Pt(10)
                p.paragraph_format.left_indent = Cm(1)
                p.paragraph_format.space_after = Pt(2)

        for q, a in sec.get("qa", []):
            p = doc.add_paragraph()
            run_q = p.add_run(f"Q：{q}")
            run_q.bold = True
            p2 = doc.add_paragraph()
            p2.add_run(f"A：{a}")
            p2.paragraph_format.space_after = Pt(8)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(FOOTER)
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    doc.save(path)
    print(f"[OK] Word 文档已生成: {path}")


# ── PDF 生成 (fpdf2) ─────────────────────────────────────

class ChinesePDF(FPDF):
    def __init__(self):
        super().__init__()
        # 注册中文字体（子集嵌入）
        self.add_font("msyh", "", "C:/Windows/Fonts/msyh.ttc")
        self.add_font("msyh", "B", "C:/Windows/Fonts/msyhbd.ttc")
        self.add_font("consolas", "", "C:/Windows/Fonts/consola.ttf")
        self.set_auto_page_break(auto=True, margin=25)

    def header(self):
        pass

    def footer(self):
        self.set_y(-15)
        self.set_font("msyh", "", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f"- {self.page_no()} -", align="C")


def make_pdf(path):
    pdf = ChinesePDF()
    pdf.set_margins(25, 20, 25)
    pdf.add_page()

    # 标题
    pdf.ln(10)
    pdf.set_font("msyh", "B", 20)
    pdf.set_text_color(26, 26, 46)
    pdf.cell(0, 14, TITLE, align="C", new_x="LMARGIN", new_y="NEXT")

    pdf.set_font("msyh", "", 12)
    pdf.set_text_color(100, 100, 100)
    pdf.cell(0, 10, SUBTITLE, align="C", new_x="LMARGIN", new_y="NEXT")

    pdf.ln(4)
    pdf.set_draw_color(200, 200, 200)
    pdf.line(25, pdf.get_y(), 185, pdf.get_y())
    pdf.ln(6)

    for sec in SECTIONS:
        # 小节标题
        pdf.set_font("msyh", "B", 13)
        pdf.set_text_color(44, 62, 80)
        pdf.cell(0, 10, sec["heading"], new_x="LMARGIN", new_y="NEXT")
        pdf.ln(1)

        # 段落
        pdf.set_font("msyh", "", 10)
        pdf.set_text_color(30, 30, 30)
        for text in sec.get("paragraphs", []):
            pdf.multi_cell(0, 6, text)
            pdf.ln(2)

        # 无序列表
        for item in sec.get("bullets", []):
            pdf.set_x(30)
            pdf.multi_cell(0, 6, f"  \u2022  {item}")
            pdf.ln(1)

        # 有序列表
        num = 0
        for item in sec.get("numbered", []):
            if item.startswith("  -"):
                pdf.set_x(40)
                pdf.multi_cell(0, 6, item)
                pdf.ln(0.5)
            else:
                num += 1
                pdf.set_x(30)
                pdf.multi_cell(0, 6, f"  {num}.  {item}")
                pdf.ln(1)

        # 表格
        tbl = sec.get("table")
        if tbl:
            headers = tbl["headers"]
            rows = tbl["rows"]
            col_count = len(headers)
            col_w = 160 / col_count

            # 表头
            pdf.set_font("msyh", "B", 9.5)
            pdf.set_fill_color(230, 240, 250)
            pdf.set_x(25)
            for h_text in headers:
                pdf.cell(col_w, 8, h_text, border=1, fill=True, align="C")
            pdf.ln()

            # 数据行
            pdf.set_font("msyh", "", 9)
            pdf.set_text_color(30, 30, 30)
            for row in rows:
                pdf.set_x(25)
                for val in row:
                    pdf.cell(col_w, 7, val, border=1, align="C")
                pdf.ln()
            pdf.ln(3)

        # 表格后的补充段落
        pdf.set_font("msyh", "", 10)
        pdf.set_text_color(30, 30, 30)
        for text in sec.get("paragraphs_after", []):
            pdf.multi_cell(0, 6, text)
            pdf.ln(2)

        # 代码块
        code_lines = sec.get("code")
        if code_lines:
            pdf.set_fill_color(245, 245, 245)
            pdf.set_font("msyh", "", 9.5)
            pdf.set_text_color(50, 50, 50)
            for line in code_lines:
                pdf.set_x(30)
                pdf.cell(150, 6, line, fill=True, new_x="LMARGIN", new_y="NEXT")
            pdf.ln(3)
            pdf.set_text_color(30, 30, 30)

        # 问答
        for q, a in sec.get("qa", []):
            pdf.set_font("msyh", "B", 10)
            pdf.set_text_color(30, 30, 30)
            pdf.multi_cell(0, 6, f"Q：{q}")
            pdf.set_font("msyh", "", 10)
            pdf.set_x(30)
            pdf.multi_cell(0, 6, f"A：{a}")
            pdf.ln(2)

        pdf.ln(3)

    # 页脚声明
    pdf.ln(5)
    pdf.set_draw_color(200, 200, 200)
    pdf.line(25, pdf.get_y(), 185, pdf.get_y())
    pdf.ln(4)
    pdf.set_font("msyh", "", 8)
    pdf.set_text_color(150, 150, 150)
    pdf.cell(0, 6, FOOTER, align="C")

    pdf.output(path)
    print(f"[OK] PDF 文档已生成: {path}")


# ── 主程序 ────────────────────────────────────────────────

if __name__ == "__main__":
    docx_path = os.path.join(OUTPUT_DIR, "Claude Code 使用说明.docx")
    pdf_path = os.path.join(OUTPUT_DIR, "Claude Code 使用说明.pdf")
    make_docx(docx_path)
    make_pdf(pdf_path)
    print("\n全部生成完毕！")
